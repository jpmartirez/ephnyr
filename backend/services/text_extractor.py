import io
from typing import List
from langchain_core.documents import Document
from pypdf import PdfReader
import docx

def extract_langchain_documents(file_bytes: bytes, file_name: str, file_type: str) -> List[Document]:
    """
    Extract raw text from PDF, DOCX, TXT, or MD file bytes into LangChain Document objects.
    """
    lower_name = file_name.lower()
    documents: List[Document] = []

    # 1. PDF Documents
    if lower_name.endswith(".pdf") or "pdf" in file_type:
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    documents.append(
                        Document(
                            page_content=text.strip(),
                            metadata={
                                "source": file_name,
                                "file_name": file_name,
                                "page": i + 1,
                                "file_type": "pdf",
                            },
                        )
                    )
            return documents
        except Exception as e:
            raise ValueError(f"Failed to parse PDF file with LangChain extractor: {str(e)}")

    # 2. DOCX Word Documents
    if lower_name.endswith(".docx") or "wordprocessingml" in file_type:
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            full_text = [p.text for p in doc.paragraphs if p.text.strip()]
            combined_text = "\n\n".join(full_text)
            if combined_text.strip():
                documents.append(
                    Document(
                        page_content=combined_text.strip(),
                        metadata={
                            "source": file_name,
                            "file_name": file_name,
                            "file_type": "docx",
                        },
                    )
                )
            return documents
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX file with LangChain extractor: {str(e)}")

    # 3. Text and Markdown Documents
    try:
        content = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        content = file_bytes.decode("latin-1", errors="ignore")

    if content.strip():
        documents.append(
            Document(
                page_content=content.strip(),
                metadata={
                    "source": file_name,
                    "file_name": file_name,
                    "file_type": "md" if lower_name.endswith(".md") else "txt",
                },
            )
        )

    return documents
